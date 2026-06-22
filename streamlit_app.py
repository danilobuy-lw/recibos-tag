"""
Gerador de Recibos de Pagamento Tag — v2 (com Autenticação Bancária)
=====================================================================
Faça upload da planilha de histórico de passagens e baixe os PDFs
preenchidos a partir do template Word "Confirmação de pagamento Tag".

Para rodar localmente:
    pip install -r requirements.txt
    streamlit run streamlit_app.py
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from datetime import datetime
from pathlib import Path

import openpyxl
import streamlit as st
from docx import Document


# ------------------------------------------------------------------
# Config
# ------------------------------------------------------------------
APP_DIR = Path(__file__).parent
TEMPLATE_PATH = APP_DIR / "template.docx"

# Mapeamento (linha no Word, label, coluna canônica do Excel)
# Tabela 0 = "Dados da passagem"
MAPPING_TABLE0 = [
    (1, "ID Concessionária", "id_concessionaria"),
    (2, "Placa",              "placa"),
    (3, "Data",               "data_passagem"),
    (4, "Concessionária",     "concessionaria"),
    (5, "Pórtico",            "praca"),
    (6, "Valor",              "taxa_pedagio"),
]
# Tabela 1 = "Dados do pagamento"
MAPPING_TABLE1 = [
    (1, "Data",              "data_pagamento"),
    (2, "Tarifa de pedágio", "valor"),
    (3, "Multa / Juros",     "multas"),
    (4, "Valor",             "valor_total"),
]
# Tabela 2 = "Autenticação Bancária"  (1 célula com a hash)
AUTH_TABLE_INDEX = 2
AUTH_CELL = (1, 0)
AUTH_COLUMN = "id_autenticacao_bancaria"

# Aliases aceitos para cada coluna canônica.
# A função normalize() já tira acentos, espaços e símbolos antes de comparar.
COLUMN_ALIASES = {
    "placa":                   ["placa"],
    "data_passagem":           ["datadapassagem", "datapassagem"],
    "concessionaria":          ["concessionaria"],
    "id_concessionaria":       ["idconcessionaria", "id"],
    "praca":                   ["praca"],
    "rodovia":                 ["rodovia"],
    "data_vencimento":         ["datadevencimento", "vencimento"],
    "valor":                   ["valor"],
    "multas":                  ["multas", "multa"],
    "juros":                   ["juros"],
    "taxa_pedagio":            ["taxapedagio"],
    "valor_total":             ["valortotal"],
    "desconto":                ["desconto"],
    "status_pagamento":        ["statuspagamento", "statusdopagamento"],
    "data_pagamento":          ["datadepagamento", "datapagamento"],
    "empresa":                 ["empresa"],
    "id_autenticacao_bancaria":["idautenticacaobancaria", "autenticacaobancaria",
                                 "idautenticacao", "autenticacao"],
}


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------
def normalize(s: str) -> str:
    if s is None: return ""
    s = str(s).strip()
    s = "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", "", s)
    return s


def map_headers(headers: tuple) -> dict:
    """{chave_canonica: indice_da_coluna}, tolerante a variações de encoding."""
    found = {}
    for idx, h in enumerate(headers):
        nh = normalize(h)
        for canonical, aliases in COLUMN_ALIASES.items():
            if nh in aliases and canonical not in found:
                found[canonical] = idx
                break

    # Prefixos para variações corrompidas (ex.: ID_CONCESSIONÝRIA -> idconcessionyria)
    PREFIX_MATCHES = {
        "id_concessionaria": "idconcession",
        "id_autenticacao_bancaria": "idautenticacao",
    }
    for idx, h in enumerate(headers):
        nh = normalize(h)
        for canonical, prefix in PREFIX_MATCHES.items():
            if canonical in found:
                continue
            if nh.startswith(prefix):
                found[canonical] = idx
                break
    return found


def set_cell_text(cell, new_text: str) -> None:
    paragraphs = cell.paragraphs
    p = paragraphs[0]
    if not p.runs:
        p.add_run(str(new_text))
    else:
        p.runs[0].text = str(new_text)
        for run in p.runs[1:]:
            run._element.getparent().remove(run._element)
    for extra_p in paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def format_placa(placa) -> str:
    if not placa: return ""
    s = str(placa).strip().replace("-", "").replace(" ", "").upper()
    return f"{s[:3]}-{s[3:]}" if len(s) == 7 else s


def fmt_datetime_passagem(value) -> str:
    """Data com hora: dd/mm/yyyy, HH:MM"""
    if value is None or value == "": return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y, %H:%M")
    s = str(value).strip()
    for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"):
        try: return datetime.strptime(s, fmt).strftime("%d/%m/%Y, %H:%M")
        except: pass
    return s


def fmt_date_pagamento(value) -> str:
    """Só data: dd/mm/yyyy"""
    if value is None or value == "": return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    s = str(value).strip()
    for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try: return datetime.strptime(s, fmt).strftime("%d/%m/%Y")
        except: pass
    return s


def fmt_money(value) -> str:
    """Sempre 'R$ XX,XX' — vírgula decimal, ponto de milhar."""
    if value is None or value == "":
        return "R$ 0,00"
    s = str(value).replace("\xa0", " ").strip()
    cleaned = s.replace("R$", "").strip().replace(".", "").replace(",", ".")
    try:
        n = float(cleaned)
        int_part, dec_part = f"{n:.2f}".split(".")
        rev = int_part[::-1]
        grouped = ".".join(rev[i:i+3] for i in range(0, len(rev), 3))[::-1]
        return f"R$ {grouped},{dec_part}"
    except ValueError:
        return s if s.startswith("R$") else f"R$ {s}"


def safe_filename(s) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", str(s)).strip("_")


def format_for_column(canonical: str, value):
    if canonical == "placa":             return format_placa(value)
    if canonical == "data_passagem":     return fmt_datetime_passagem(value)
    if canonical == "data_pagamento":    return fmt_date_pagamento(value)
    if canonical in ("valor", "taxa_pedagio", "multas", "juros", "valor_total"):
        return fmt_money(value)
    if value is None:                    return ""
    return str(value)


# ------------------------------------------------------------------
# Geração
# ------------------------------------------------------------------
def gerar_um_recibo(template_path: Path, row, hidx, idx: int) -> tuple[str, bytes]:
    doc = Document(str(template_path))
    tables = doc.tables

    for row_idx, label, canonical in MAPPING_TABLE0:
        raw = row[hidx[canonical]] if canonical in hidx else None
        set_cell_text(tables[0].rows[row_idx].cells[1], format_for_column(canonical, raw))

    for row_idx, label, canonical in MAPPING_TABLE1:
        raw = row[hidx[canonical]] if canonical in hidx else None
        set_cell_text(tables[1].rows[row_idx].cells[1], format_for_column(canonical, raw))

    # Tabela 2: Autenticação Bancária (1 célula)
    if len(tables) > AUTH_TABLE_INDEX and AUTH_COLUMN in hidx:
        auth_val = row[hidx[AUTH_COLUMN]]
        auth_cell = tables[AUTH_TABLE_INDEX].rows[AUTH_CELL[0]].cells[AUTH_CELL[1]]
        set_cell_text(auth_cell, str(auth_val) if auth_val is not None else "")

    id_v   = row[hidx["id_concessionaria"]] if "id_concessionaria" in hidx else f"linha{idx}"
    placa_v = format_placa(row[hidx["placa"]]) if "placa" in hidx else ""
    fname = safe_filename(f"Recibo_{idx:02d}_{id_v}_{placa_v}")

    buf = io.BytesIO()
    doc.save(buf)
    return fname + ".docx", buf.getvalue()


def converter_batch_para_pdf(docx_paths: list[Path], outdir: Path) -> dict[str, bytes]:
    """Converte vários docx para PDF em UMA chamada do LibreOffice."""
    if not docx_paths:
        return {}
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(outdir)] + [str(p) for p in docx_paths],
            capture_output=True, timeout=600,
        )
    except FileNotFoundError:
        return {}
    except Exception:
        return {}

    out = {}
    for p in docx_paths:
        pdf_path = outdir / (p.stem + ".pdf")
        if pdf_path.exists():
            out[p.stem] = pdf_path.read_bytes()
            try: pdf_path.unlink()
            except: pass
    for f in outdir.glob("*.tmp"):
        try: f.unlink()
        except: pass
    for f in outdir.glob(".~lock.*"):
        try: f.unlink()
        except: pass
    return out


def processar_planilha(file_bytes: bytes, somente_pago: bool, batch_size: int = 50,
                       progress_cb=None):
    """Lê planilha, gera os recibos em lotes de `batch_size` para evitar estouro."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("A planilha está vazia.")
    headers = rows[0]
    data_rows = rows[1:]
    hidx = map_headers(headers)

    required = {"placa", "data_passagem", "concessionaria", "id_concessionaria",
                "praca", "valor", "taxa_pedagio", "valor_total",
                "data_pagamento", "id_autenticacao_bancaria"}
    missing = required - set(hidx.keys())
    if missing:
        nice = {
            "placa": "Placa",
            "data_passagem": "Data da Passagem",
            "concessionaria": "Concessionária",
            "id_concessionaria": "ID_CONCESSIONÁRIA",
            "praca": "Praça",
            "valor": "Valor",
            "taxa_pedagio": "Taxa pedágio",
            "valor_total": "Valor total",
            "data_pagamento": "Data de pagamento",
            "id_autenticacao_bancaria": "ID AUTENTICAÇÃO BANCÁRIA",
        }
        raise ValueError(
            "A planilha não tem as colunas: " + ", ".join(nice[m] for m in missing)
        )

    if somente_pago and "status_pagamento" in hidx:
        idx_s = hidx["status_pagamento"]
        target = [r for r in data_rows if r[idx_s] and normalize(r[idx_s]) == "pago"]
    else:
        target = list(data_rows)

    if not target:
        return []

    workdir = Path(tempfile.mkdtemp(prefix="recibos_"))
    results = []
    try:
        # Gera todos os docx
        all_specs = []
        for i, row in enumerate(target, start=1):
            fname_docx, docx_bytes = gerar_um_recibo(TEMPLATE_PATH, row, hidx, i)
            p = workdir / fname_docx
            p.write_bytes(docx_bytes)
            all_specs.append((fname_docx, p))
            if progress_cb:
                progress_cb(i, len(target) * 2)

        # Converte em LOTES (evita estouro de tempo do LibreOffice)
        pdf_map = {}
        for start in range(0, len(all_specs), batch_size):
            chunk = all_specs[start:start + batch_size]
            paths = [p for _, p in chunk]
            converted = converter_batch_para_pdf(paths, workdir)
            pdf_map.update(converted)
            if progress_cb:
                done = start + len(chunk)
                progress_cb(len(target) + done, len(target) * 2)

        # Monta resultados
        for fname_docx, p in all_specs:
            stem = p.stem
            if stem in pdf_map:
                results.append((stem + ".pdf", pdf_map[stem], "pdf"))
            else:
                results.append((fname_docx, p.read_bytes(), "docx"))

        return results
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def zip_files(files: list[tuple[str, bytes, str]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data, _ in files:
            zf.writestr(name, data)
    return buf.getvalue()


# ------------------------------------------------------------------
# UI
# ------------------------------------------------------------------
st.set_page_config(page_title="Gerador de Recibos de Pagamento Tag",
                   page_icon="🧾", layout="centered")

st.title("🧾 Gerador de Recibos de Pagamento Tag")
st.markdown(
    "Envie a planilha de **histórico de passagens** (`.xlsx`). "
    "O app preenche o template Word *Confirmação de pagamento Tag* "
    "e gera um PDF por linha."
)

with st.expander("ℹ️ Quais colunas a planilha precisa ter?"):
    st.markdown("""
- **Placa**
- **Data da Passagem**
- **Concessionária**
- **ID_CONCESSIONÁRIA**
- **Praça**
- **Valor**
- **Taxa pedágio**
- **Valor total**
- **Data de pagamento**
- **ID AUTENTICAÇÃO BANCÁRIA**
- **Status pagamento** *(opcional — usado para filtrar)*
- **Multas** *(opcional — vira "Multa / Juros" no recibo)*
""")

uploaded = st.file_uploader("Planilha de histórico (.xlsx)", type=["xlsx"])

somente_pago = st.checkbox(
    "Gerar apenas para linhas com Status Pagamento = Pago",
    value=True,
)

batch_size = st.number_input(
    "Tamanho do lote de conversão (PDFs por vez)",
    min_value=10, max_value=200, value=50, step=10,
    help="Lotes menores são mais seguros em servidores com pouca memória.",
)

if uploaded is not None:
    file_bytes = uploaded.read()
    try:
        wb_prev = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws_prev = wb_prev.active
        prev_rows = list(ws_prev.iter_rows(values_only=True))
        n_total = max(0, len(prev_rows) - 1)
        headers = prev_rows[0] if prev_rows else ()
        hidx = map_headers(headers)
        if somente_pago and "status_pagamento" in hidx:
            n_alvo = sum(
                1 for r in prev_rows[1:]
                if r[hidx["status_pagamento"]]
                and normalize(r[hidx["status_pagamento"]]) == "pago"
            )
        else:
            n_alvo = n_total

        c1, c2 = st.columns(2)
        c1.metric("Linhas na planilha", n_total)
        c2.metric("Recibos a gerar", n_alvo)
    except Exception as e:
        st.error(f"Erro ao ler a planilha: {e}")
        st.stop()

    if st.button("🚀 Gerar Recibos", type="primary", disabled=(n_alvo == 0)):
        progress = st.progress(0, text="Iniciando...")

        def cb(i, total):
            progress.progress(min(i / total, 1.0),
                              text=f"Processando {i} de {total} passos...")

        try:
            with st.spinner("Processando..."):
                results = processar_planilha(
                    file_bytes, somente_pago, batch_size, cb
                )
            progress.progress(1.0, text="Concluído!")

            n_pdf = sum(1 for _, _, kind in results if kind == "pdf")
            n_docx = sum(1 for _, _, kind in results if kind == "docx")

            if n_pdf and not n_docx:
                st.success(f"✅ {n_pdf} PDF(s) gerados!")
            elif n_docx and not n_pdf:
                st.warning(
                    f"⚠️ {n_docx} arquivo(s) gerados em .docx — "
                    "LibreOffice/Word indisponível no servidor."
                )
            else:
                st.info(
                    f"Gerados {n_pdf} PDF(s) e {n_docx} .docx. "
                    "Os .docx ficaram porque o conversor falhou neles. "
                    "Tente baixar o ZIP completo, ou reduza o tamanho do lote."
                )

            zip_bytes = zip_files(results)
            st.download_button(
                "⬇️ Baixar todos os recibos (.zip)",
                data=zip_bytes,
                file_name="recibos_tag.zip",
                mime="application/zip",
                type="primary",
            )

            with st.expander("Baixar individualmente"):
                for name, data, kind in results:
                    mime = "application/pdf" if kind == "pdf" else \
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.download_button(
                        f"⬇️ {name}",
                        data=data,
                        file_name=name,
                        mime=mime,
                        key=f"dl_{name}",
                    )
        except Exception as e:
            st.error(f"Erro: {e}")
            st.exception(e)

st.markdown("---")
st.caption("v2 — com Autenticação Bancária • Valores em R$ XX,XX • Conversão em lote")
