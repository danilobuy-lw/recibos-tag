"""
Gerador de Recibos de Pagamento Tag — app web (Streamlit)
=========================================================
Faça upload da planilha de histórico de passagens e baixe os PDFs
preenchidos a partir do template Word "Confirmação de pagamento Tag".

Para rodar localmente:
    pip install -r requirements.txt
    streamlit run streamlit_app.py

Para publicar: ver README.md (Streamlit Community Cloud).
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from datetime import datetime
from pathlib import Path

import openpyxl
import streamlit as st
from docx import Document


# ------------------------------------------------------------------
# Config
# ------------------------------------------------------------------
APP_DIR = Path(__file__).parent
TEMPLATE_PATH = APP_DIR / "template.docx"

# Mapeamento: (índice da linha no Word, label, chave canônica da coluna)
MAPPING_TABLE0 = [
    (1, "ID",             "id_concessionaria"),
    (2, "Placa",          "placa"),
    (3, "Data",           "data_passagem"),
    (4, "Concessionária", "praca"),
    (5, "Pórtico",        "rodovia"),
    (6, "Valor",          "valor"),
]
MAPPING_TABLE1 = [
    (1, "Data",              "data_pagamento"),
    (2, "Tarifa de pedágio", "taxa_pedagio"),
    (5, "Valor",             "taxa_pedagio"),
]

# Aliases aceitos por coluna canônica (em minúsculas, sem acento e sem espaços)
COLUMN_ALIASES = {
    "placa":               ["placa"],
    "data_passagem":       ["datadapassagem", "datapassagem"],
    "praca":               ["praca"],
    "rodovia":             ["rodovia"],
    "valor":               ["valor"],
    "taxa_pedagio":        ["taxapedagio"],
    "status_pagamento":    ["statuspagamento", "statusdopagamento"],
    "data_pagamento":      ["datapagamento", "datadepagamento"],
    "id_concessionaria":   ["idconcessionaria", "id"],
}


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------
def normalize(s: str) -> str:
    """Normaliza string para comparação: minúscula, sem acento, sem espaços/símbolos."""
    if s is None:
        return ""
    s = str(s).strip()
    # Remove acentos
    s = "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", "", s)
    return s


def map_headers(headers: tuple) -> dict:
    """Retorna {chave_canonica: indice_da_coluna} aceitando variações de nome.

    Faz duas passadas:
      1) match exato contra aliases
      2) match por prefixo (tolera erros de encoding, ex.: ID_CONCESSIONÝRIA)
    """
    found = {}
    # Pass 1: alias exato
    for idx, h in enumerate(headers):
        nh = normalize(h)
        for canonical, aliases in COLUMN_ALIASES.items():
            if nh in aliases and canonical not in found:
                found[canonical] = idx
                break

    # Pass 2: prefixos especiais para casos com encoding corrompido
    PREFIX_MATCHES = {
        "id_concessionaria": "idconcession",  # cobre idconcessionyria, idconcessionaria...
    }
    for idx, h in enumerate(headers):
        nh = normalize(h)
        for canonical, prefix in PREFIX_MATCHES.items():
            if canonical in found:
                continue
            if nh.startswith(prefix):
                found[canonical] = idx
                break

    return found


def set_cell_text(cell, new_text: str) -> None:
    paragraphs = cell.paragraphs
    p = paragraphs[0]
    if not p.runs:
        p.add_run(str(new_text))
    else:
        first_run = p.runs[0]
        first_run.text = str(new_text)
        for run in p.runs[1:]:
            run._element.getparent().remove(run._element)
    for extra_p in paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def format_placa(placa) -> str:
    if not placa:
        return ""
    s = str(placa).strip().replace("-", "").replace(" ", "").upper()
    return f"{s[:3]}-{s[3:]}" if len(s) == 7 else s


def fmt_datetime_passagem(value) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y, %H:%M")
    s = str(value).strip()
    for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(s, fmt).strftime("%d/%m/%Y, %H:%M")
        except ValueError:
            continue
    return s


def fmt_date_pagamento(value) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    s = str(value).strip()
    for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return s


def clean_currency(value) -> str:
    if value is None or value == "":
        return "R$ 0,00"
    s = str(value).replace("\xa0", " ").strip()
    if not s.startswith("R$"):
        s = f"R$ {s}"
    return re.sub(r"R\$\s*", "R$ ", s)


def safe_filename(s) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", str(s)).strip("_")


def format_for_column(canonical: str, value):
    if canonical == "placa":            return format_placa(value)
    if canonical == "data_passagem":    return fmt_datetime_passagem(value)
    if canonical == "data_pagamento":   return fmt_date_pagamento(value)
    if canonical in ("valor", "taxa_pedagio"): return clean_currency(value)
    if value is None: return ""
    return str(value)


# ------------------------------------------------------------------
# Geração
# ------------------------------------------------------------------
def gerar_um_recibo(template_path: Path, row, hidx, idx: int) -> tuple[str, bytes]:
    """Gera um docx preenchido em memória. Retorna (nome_arquivo, bytes_docx)."""
    doc = Document(str(template_path))
    tables = doc.tables

    for row_idx, label, canonical in MAPPING_TABLE0:
        raw = row[hidx[canonical]] if canonical in hidx else None
        set_cell_text(tables[0].rows[row_idx].cells[1], format_for_column(canonical, raw))

    for row_idx, label, canonical in MAPPING_TABLE1:
        raw = row[hidx[canonical]] if canonical in hidx else None
        set_cell_text(tables[1].rows[row_idx].cells[1], format_for_column(canonical, raw))

    id_v   = row[hidx["id_concessionaria"]] if "id_concessionaria" in hidx else f"linha{idx}"
    placa_v = format_placa(row[hidx["placa"]]) if "placa" in hidx else ""
    fname = safe_filename(f"Recibo_{idx:02d}_{id_v}_{placa_v}")

    buf = io.BytesIO()
    doc.save(buf)
    return fname + ".docx", buf.getvalue()


def converter_batch_para_pdf(docx_paths: list[Path], outdir: Path) -> dict[str, bytes]:
    """Converte vários docx para PDF numa única chamada do LibreOffice.

    Retorna {stem: pdf_bytes} para os arquivos convertidos. Faltantes = falhou.
    """
    if not docx_paths:
        return {}
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(outdir)] + [str(p) for p in docx_paths],
            capture_output=True, timeout=600,
        )
    except FileNotFoundError:
        return {}
    except Exception:
        return {}

    out = {}
    for p in docx_paths:
        pdf_path = outdir / (p.stem + ".pdf")
        if pdf_path.exists():
            out[p.stem] = pdf_path.read_bytes()
            try: pdf_path.unlink()
            except: pass
    # Limpa lixo do LibreOffice
    for f in outdir.glob("*.tmp"):
        try: f.unlink()
        except: pass
    for f in outdir.glob(".~lock.*"):
        try: f.unlink()
        except: pass
    return out


def processar_planilha(file_bytes: bytes, somente_pago: bool, progress_cb=None):
    """Lê a planilha e gera os arquivos. Retorna lista [(nome.pdf, bytes), ...]."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("A planilha está vazia.")
    headers = rows[0]
    data_rows = rows[1:]
    hidx = map_headers(headers)

    required = {"placa", "data_passagem", "praca", "rodovia", "valor",
                "taxa_pedagio", "data_pagamento", "id_concessionaria"}
    missing = required - set(hidx.keys())
    if missing:
        nice = {
            "placa": "Placa", "data_passagem": "Data da Passagem",
            "praca": "Praça", "rodovia": "Rodovia", "valor": "Valor",
            "taxa_pedagio": "Taxa pedágio", "data_pagamento": "Data Pagamento",
            "id_concessionaria": "ID_CONCESSIONÁRIA",
        }
        raise ValueError("A planilha não tem as colunas: "
                         + ", ".join(nice[m] for m in missing))

    if somente_pago and "status_pagamento" in hidx:
        idx_s = hidx["status_pagamento"]
        target = [r for r in data_rows
                  if r[idx_s] and normalize(r[idx_s]) == "pago"]
    else:
        target = list(data_rows)

    workdir = Path(tempfile.mkdtemp(prefix="recibos_"))
    docx_specs = []  # list of (fname_docx, path, idx)
    try:
        # 1) Gera todos os docx no disco
        for i, row in enumerate(target, start=1):
            fname_docx, docx_bytes = gerar_um_recibo(TEMPLATE_PATH, row, hidx, i)
            p = workdir / fname_docx
            p.write_bytes(docx_bytes)
            docx_specs.append((fname_docx, p, i))
            if progress_cb:
                progress_cb(i, len(target) * 2)  # 1ª metade da barra

        # 2) Converte tudo em lote
        pdfs = converter_batch_para_pdf([p for _, p, _ in docx_specs], workdir)

        # 3) Monta resultado
        results = []
        for j, (fname_docx, p, _) in enumerate(docx_specs, start=1):
            stem = p.stem
            if stem in pdfs:
                results.append((stem + ".pdf", pdfs[stem], "pdf"))
            else:
                # fallback: devolve o docx
                results.append((fname_docx, p.read_bytes(), "docx"))
            if progress_cb:
                progress_cb(len(target) + j, len(target) * 2)

        return results
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def zip_files(files: list[tuple[str, bytes, str]]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data, _ in files:
            zf.writestr(name, data)
    return buf.getvalue()


# ------------------------------------------------------------------
# UI
# ------------------------------------------------------------------
st.set_page_config(page_title="Gerador de Recibos de Pagamento Tag",
                   page_icon="🧾", layout="centered")

st.title("🧾 Gerador de Recibos de Pagamento Tag")
st.markdown(
    "Envie a planilha de **histórico de passagens** (`.xlsx`). "
    "O app preenche o template Word *Confirmação de pagamento Tag* "
    "e gera um PDF por linha."
)

with st.expander("ℹ️ Quais colunas a planilha precisa ter?"):
    st.markdown("""
- **Placa**
- **Data da Passagem**
- **Praça**
- **Rodovia**
- **ID_CONCESSIONÁRIA** (aceita também `ID`)
- **Valor**
- **Taxa pedágio**
- **Status pagamento** (opcional — usado para filtrar)
- **Data Pagamento** (ou *Data de pagamento*)
""")

uploaded = st.file_uploader("Planilha de histórico (.xlsx)", type=["xlsx"])

somente_pago = st.checkbox(
    "Gerar apenas para linhas com Status Pagamento = Pago",
    value=True,
)

if uploaded is not None:
    file_bytes = uploaded.read()

    # Preview rápido
    try:
        wb_prev = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws_prev = wb_prev.active
        prev_rows = list(ws_prev.iter_rows(values_only=True))
        n_total = max(0, len(prev_rows) - 1)
        headers = prev_rows[0] if prev_rows else ()
        hidx = map_headers(headers)
        if somente_pago and "status_pagamento" in hidx:
            n_alvo = sum(
                1 for r in prev_rows[1:]
                if r[hidx["status_pagamento"]]
                and normalize(r[hidx["status_pagamento"]]) == "pago"
            )
        else:
            n_alvo = n_total

        c1, c2 = st.columns(2)
        c1.metric("Linhas na planilha", n_total)
        c2.metric("Recibos a gerar", n_alvo)
    except Exception as e:
        st.error(f"Erro ao ler a planilha: {e}")
        st.stop()

    if st.button("🚀 Gerar Recibos", type="primary", disabled=(n_alvo == 0)):
        progress = st.progress(0, text="Iniciando...")

        def cb(i, total):
            progress.progress(i / total, text=f"Gerando recibo {i} de {total}...")

        try:
            with st.spinner("Processando..."):
                results = processar_planilha(file_bytes, somente_pago, cb)
            progress.progress(1.0, text="Concluído!")

            n_pdf = sum(1 for _, _, kind in results if kind == "pdf")
            n_docx = sum(1 for _, _, kind in results if kind == "docx")

            if n_pdf and not n_docx:
                st.success(f"✅ {n_pdf} PDF(s) gerados!")
            elif n_docx and not n_pdf:
                st.warning(
                    f"⚠️ {n_docx} arquivo(s) gerados em .docx — "
                    "LibreOffice/Word não está disponível neste servidor."
                )
            else:
                st.info(f"Gerados {n_pdf} PDF(s) e {n_docx} .docx.")

            zip_bytes = zip_files(results)
            st.download_button(
                "⬇️ Baixar todos os recibos (.zip)",
                data=zip_bytes,
                file_name="recibos_tag.zip",
                mime="application/zip",
                type="primary",
            )

            with st.expander("Baixar individualmente"):
                for name, data, kind in results:
                    mime = "application/pdf" if kind == "pdf" else \
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.download_button(
                        f"⬇️ {name}",
                        data=data,
                        file_name=name,
                        mime=mime,
                        key=f"dl_{name}",
                    )
        except Exception as e:
            st.error(f"Erro: {e}")
            st.exception(e)

st.markdown("---")
st.caption("Template embutido: *Confirmação de pagamento Tag.docx*")
